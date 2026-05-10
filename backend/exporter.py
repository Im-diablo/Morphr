"""
exporter.py — Generate PDF (fpdf2) and DOCX (python-docx) from AI-edited LaTeX.

PDF:  fpdf2 fallback when Tectonic is unavailable.
DOCX: always generated; parses the full LaTeX document structure so the
      output matches the compiled PDF as closely as possible.
"""

import re
import os
import logging

logger = logging.getLogger(__name__)

# ── LaTeX stripping ───────────────────────────────────────────────────────────

_COMMENT_RE = re.compile(r"%%.*$|%(?!%).*$", re.MULTILINE)


def strip_latex(text: str) -> str:
    """Strip LaTeX markup and return readable plain text."""
    text = _COMMENT_RE.sub("", text)
    text = re.sub(r"\\(?:textbf|textit|emph|underline|texttt|textrm|textsc)\{([^}]*)\}", r"\1", text)
    text = re.sub(r"\\href\{[^}]*\}\{([^}]*)\}", r"\1", text)
    text = re.sub(r"\\url\{([^}]*)\}", r"\1", text)
    text = re.sub(r"\\section\*?\s*\{[^}]*\}", "", text)
    text = re.sub(r"\\subsection\*?\s*\{[^}]*\}", "", text)
    text = re.sub(r"\\[a-zA-Z]+\*?(?:\[[^\]]*\])?(?:\{[^}]*\})*", " ", text)
    text = re.sub(r"[{}\$]", "", text)
    text = text.replace("|", " | ").replace("\\\\", "\n")
    text = re.sub(r"[ \t]{2,}", " ", text)
    text = re.sub(r"\n{3,}", "\n\n", text)
    return text.strip()


# ── Full-document section parser ──────────────────────────────────────────────

def _parse_document(tex_content: str) -> list:
    """
    Parse a complete LaTeX resume into a list of (heading, content) tuples.

    Strategy:
    1. Extract the \\begin{center}...\\end{center} header block.
    2. Split the rest by \\section*{Name} or \\section{Name}.
    3. For each block, strip LaTeX and keep meaningful lines.

    Returns: [(heading_str, content_str), ...]
    where heading_str is "" for the contact header.
    """
    results = []

    # 1. Header block (name + contact)
    center = re.search(r"\\begin\{center\}(.*?)\\end\{center\}", tex_content, re.DOTALL)
    if center:
        header = strip_latex(center.group(1))
        if header.strip():
            results.append(("", header.strip()))

    # 2. Body after \begin{document} (or whole doc if marker absent)
    body_start = tex_content.find(r"\begin{document}")
    body = tex_content[body_start:] if body_start != -1 else tex_content

    # Split on \section*{...} or \section{...}
    sec_pattern = re.compile(r"\\section\*?\s*\{([^}]+)\}", re.IGNORECASE)
    parts = sec_pattern.split(body)
    # parts = [preamble, name1, content1, name2, content2, ...]
    i = 1
    while i + 1 < len(parts):
        name = parts[i].strip()
        content_raw = parts[i + 1]
        # Remove nested \subsection, itemize env tags, etc.
        content = strip_latex(content_raw).strip()
        if content:
            results.append((name, content))
        i += 2

    # 3. Fallback: if nothing found, just strip everything
    if not results:
        results.append(("Resume", strip_latex(tex_content)))

    return results


# ── EDITABLE-section extractor (for fpdf2 fallback only) ─────────────────────

_KNOWN_LABELS = {
    "SUMMARY": "Profile Summary", "SKILLS": "Skills",
    "EXPERIENCE": "Experience",   "PROJECTS": "Projects",
    "CONTENT": "Resume",
}


def _extract_sections(tex_content: str) -> list:
    """
    Parse %% EDITABLE:X / %% END:X blocks. Returns (key, label, text) tuples.
    Handles both named (SUMMARY) and numeric (1, 2, 3) keys.
    """
    pattern = re.compile(r"%%\s*EDITABLE:(\w+)\s*\n(.*?)%%\s*END:\1",
                         re.DOTALL | re.IGNORECASE)
    results = []
    for m in pattern.finditer(tex_content):
        key = m.group(1).upper()
        raw = m.group(2)
        # Derive label
        if key in _KNOWN_LABELS:
            label = _KNOWN_LABELS[key]
        else:
            hm = re.search(r"\\section\*?\s*\{([^}]+)\}", raw)
            label = hm.group(1).strip() if hm else (f"Section {key}" if key.isdigit() else key.capitalize())
        content = strip_latex(raw).strip()
        if content:
            results.append((key, label, content))

    if not results:
        results.append(("CONTENT", "Resume", strip_latex(tex_content)))
    return results


# ── latin-1 safe helper ───────────────────────────────────────────────────────

def _safe(text: str) -> str:
    subs = {"\u2014": "-", "\u2013": "-", "\u2019": "'", "\u2018": "'",
            "\u201c": '"', "\u201d": '"', "\u2022": "-", "\u00a0": " ",
            "\u2026": "...", "\u00b7": "-"}
    for ch, r in subs.items():
        text = text.replace(ch, r)
    return text.encode("latin-1", errors="replace").decode("latin-1")


# ── PDF via fpdf2 (fallback only) ────────────────────────────────────────────

def generate_pdf(tex_content: str, output_path: str) -> str:
    """Generate a formatted PDF from LaTeX using fpdf2. Returns output_path."""
    from fpdf import FPDF

    sections = _extract_sections(tex_content)
    pdf = FPDF(format="A4")
    pdf.set_auto_page_break(auto=True, margin=18)
    pdf.add_page()
    pdf.set_margins(20, 20, 20)

    pdf.set_text_color(255, 215, 0)
    pdf.set_font("Helvetica", "B", 20)
    pdf.cell(0, 12, "Resume", ln=True, align="C")
    pdf.set_text_color(160, 160, 160)
    pdf.set_font("Helvetica", "", 8)
    pdf.cell(0, 5, "Generated by Morphr - AI-Powered Resume Tool", ln=True, align="C")
    pdf.ln(4)
    pdf.set_draw_color(255, 215, 0)
    pdf.set_line_width(0.4)
    pdf.line(20, pdf.get_y(), 190, pdf.get_y())
    pdf.ln(6)

    for _key, label, content in sections:
        pdf.set_text_color(255, 215, 0)
        pdf.set_font("Helvetica", "B", 10)
        pdf.set_x(pdf.l_margin)
        pdf.cell(0, 7, _safe(label.upper()), ln=True)
        y = pdf.get_y()
        pdf.set_draw_color(255, 215, 0)
        pdf.set_line_width(0.2)
        pdf.line(20, y, 190, y)
        pdf.ln(3)

        pdf.set_text_color(30, 30, 30)
        pdf.set_font("Helvetica", "", 9)
        for line in content.split("\n"):
            line = _safe(line.strip())
            if not line:
                pdf.ln(2)
                continue
            pdf.set_x(pdf.l_margin)
            pdf.multi_cell(0, 5, line, align="L")
        pdf.ln(4)

    pdf.output(output_path)
    logger.info("fpdf2 PDF: %s (%d bytes)", output_path, os.path.getsize(output_path))
    return output_path


# ── DOCX via python-docx (full-document parse) ───────────────────────────────

def generate_docx(tex_content: str, output_path: str) -> str:
    """
    Generate a DOCX that mirrors the full LaTeX resume structure.
    Parses ALL \\section* blocks (not just EDITABLE ones).
    Returns output_path.
    """
    from docx import Document
    from docx.shared import Pt, RGBColor, Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    doc = Document()

    for sec in doc.sections:
        sec.top_margin    = Inches(0.75)
        sec.bottom_margin = Inches(0.75)
        sec.left_margin   = Inches(0.9)
        sec.right_margin  = Inches(0.9)

    parsed = _parse_document(tex_content)

    for heading, content in parsed:
        if heading == "":
            # Contact header — centered
            for line in content.split("\n"):
                line = line.strip()
                if not line:
                    continue
                # First non-empty line = name
                if not any(p.text for p in doc.paragraphs):
                    p = doc.add_paragraph()
                    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    run = p.add_run(line)
                    run.bold = True
                    run.font.size = Pt(16)
                    run.font.color.rgb = RGBColor(0x0A, 0x0A, 0x14)
                else:
                    p = doc.add_paragraph()
                    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    run = p.add_run(line)
                    run.font.size = Pt(9)
                    run.font.color.rgb = RGBColor(80, 80, 80)
        else:
            # Section heading
            h = doc.add_heading(heading, level=1)
            if h.runs:
                h.runs[0].font.color.rgb = RGBColor(0x1A, 0x1A, 0x2E)
                h.runs[0].font.size = Pt(11)

            # Section body
            for line in content.split("\n"):
                line = line.strip()
                if not line:
                    continue
                is_bullet = line.startswith(("-", "*", "·", "\u2022"))
                if is_bullet:
                    text = line.lstrip("-*·\u2022 ").strip()
                    try:
                        p = doc.add_paragraph(text, style="List Bullet")
                    except KeyError:
                        p = doc.add_paragraph("• " + text)
                else:
                    p = doc.add_paragraph(line)
                if p.runs:
                    p.runs[0].font.size = Pt(10)

    doc.save(output_path)
    logger.info("DOCX: %s (%d bytes, %d sections)", output_path,
                os.path.getsize(output_path), len(parsed))
    return output_path
